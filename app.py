import difflib
import os
import re
from openai import OpenAI
from docx import Document
from flask import Flask, render_template, request
from pdf2docx import parse
from unidecode import unidecode
from pdf2image import convert_from_path
import pytesseract
from markdown import markdown
from docx import Document
from bs4 import BeautifulSoup
from dotenv import load_dotenv

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads/'
load_dotenv()

OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
OPENAI_API_URL = os.getenv('OPENAI_API_URL')


if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

def convert_to_plain_text(text):
    text_no_punct = re.sub(r'[^\w\s]', '', text)
    return unidecode(text_no_punct)

def read_docx_with_positions(file):
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def compare_and_highlight_with_positions(text1, text2):
    text1_plain = convert_to_plain_text(text1)
    text2_plain = convert_to_plain_text(text2)

    words1 = text1_plain.split()
    words2 = text2_plain.split()

    original_words = text1.split()
    compared_words = text2.split()

    diff = list(difflib.ndiff(words1, words2))

    highlighted_text1 = []
    highlighted_text2 = []

    i1, i2 = 0, 0

    for change in diff:
        if change.startswith("- "):
            original_word = original_words[i1] if i1 < len(original_words) else ""
            highlighted_text1.append(f'<span style="color:red; background-color:pink">{original_word}</span>')
            i1 += 1
        elif change.startswith("+ "):
            new_word = compared_words[i2] if i2 < len(compared_words) else ""
            highlighted_text2.append(f'<span style="color:green; background-color:lightgreen">{new_word}</span>')
            i2 += 1
        else:
            if i1 < len(original_words):
                highlighted_text1.append(original_words[i1])
            if i2 < len(compared_words):
                highlighted_text2.append(compared_words[i2])
            i1 += 1
            i2 += 1

    return ' '.join(highlighted_text1), ' '.join(highlighted_text2)

def pdf_to_text(pdf_file):
    images = convert_from_path(pdf_file)
    pages_text = []
    for img in images:
        text = pytesseract.image_to_string(img, lang='vie')
        pages_text.append(text)
    return pages_text

def markdown_to_docx(md_text, output_filename):
    html = markdown(md_text)
    doc = Document()
    soup = BeautifulSoup(html, 'html.parser')

    for element in soup.descendants:
        if element.name in ['p', 'h1', 'h2', 'h3', 'li']:
            doc.add_paragraph(element.get_text())

    doc.save(output_filename)

def parse_ocr(pdf_file, output_docx_file):
    parse_text = pdf_to_text(pdf_file)
    if parse_text:
        content_md = normalize_res(result=parse_text)
        if content_md:
            markdown_to_docx(content_md, output_docx_file)

def normalize_res(result):
    openai = OpenAI(
        api_key="KKu4peEZRhj4ndElpo5K6SX6ImeP4qLs",
        base_url="https://api.deepinfra.com/v1/openai",
    )

    chat_completion = openai.chat.completions.create(
        model="meta-llama/Meta-Llama-3.1-405B-Instruct",
        messages=[{
            "role": "user",
            "content": f"""
                    You are an expert in Vietnamese text normalization. Please help me normalize the text according to the prompt below:

                    ---

                    **Prompt:**

                    "Please normalize the following text according to the following rules:  
                    1. Reformat the text to make it more readable and clearer.  
                    2. Correct all typos and spelling mistakes according to standard Vietnamese language.  
                    3. Ensure that important information such as contract numbers, addresses, phone numbers, tax IDs, and other identifying details are preserved.  
                    4. Adjust sentence structure, punctuation, and spacing between words to ensure grammatical accuracy.  
                    5. For monetary values, standardize the format (e.g., 10.560.000 VND) while keeping the numerical values unchanged.  
                    6. Ensure proper usage of symbols like hyphens, periods, and commas.  
                    7. Remove any unnecessary characters or unrelated symbols (such as printing errors).  
                    8. If any parts of the text are missing or incomplete, simply ignore them without attempting to infer their meaning.

                    **Text:**

                    {".".join(result)}

                    **Desired output:**  
                    Return the normalized and easy-to-read text according to the above rules."
                    """
        }],
    )

    return chat_completion.choices[0].message.content

def convert_pdf_to_docx(pdf_file, output_docx_file):
    parse(pdf_file, output_docx_file)
    doc = Document(output_docx_file)
    if not doc.paragraphs or all(not para.text.strip() for para in doc.paragraphs):
        parse_ocr(pdf_file, output_docx_file)
    return output_docx_file

@app.route('/', methods=['GET', 'POST'])
def upload_and_compare():
    if request.method == 'POST':
        docx_file = request.files['file1']
        file2 = request.files['file2']

        docx_file_path = os.path.join(app.config['UPLOAD_FOLDER'], docx_file.filename)
        file2_path = os.path.join(app.config['UPLOAD_FOLDER'], file2.filename)
        docx_file.save(docx_file_path)
        file2.save(file2_path)

        file2_extension = os.path.splitext(file2.filename)[1].lower()

        if file2_extension == '.pdf':
            converted_docx_file = os.path.join(app.config['UPLOAD_FOLDER'], 'converted_file.docx')
            convert_pdf_to_docx(file2_path, converted_docx_file)
            comparison_text = read_docx_with_positions(converted_docx_file)
        else:
            comparison_text = read_docx_with_positions(file2_path)

        original_text = read_docx_with_positions(docx_file_path)
        highlighted_text1, highlighted_text2 = compare_and_highlight_with_positions(original_text, comparison_text)

        return render_template('compare.html', text1=highlighted_text1, text2=highlighted_text2)

    return render_template('compare.html', text1=None, text2=None)

if __name__ == "__main__":
    app.run(debug=True)
